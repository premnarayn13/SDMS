"""
Text Extraction Service
Extracts text from PDF, DOCX, TXT files for indexing and search.
"""
import io
import logging
from typing import Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various document formats"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.txt', '.md'}
    
    def extract(self, file_content: bytes, filename: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract text from file.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        ext = Path(filename).suffix.lower()
        
        if ext not in self.supported_formats:
            return None, f"Unsupported format: {ext}"
        
        try:
            if ext == '.pdf':
                return self._extract_pdf(file_content), None
            elif ext == '.docx':
                return self._extract_docx(file_content), None
            elif ext in {'.txt', '.md'}:
                return self._extract_text(file_content), None
            else:
                return None, f"Handler not implemented for {ext}"
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}")
            return None, str(e)
    
    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF using pdfplumber"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            logger.warning("pdfplumber not installed, PDF extraction disabled")
            return ""
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return ""
    
    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            logger.warning("python-docx not installed, DOCX extraction disabled")
            return ""
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return ""
    
    def _extract_text(self, content: bytes) -> str:
        """Extract text from plain text files"""
        try:
            # Try UTF-8 first
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                return content.decode('latin-1')
            except Exception as e:
                logger.error(f"Text decoding error: {str(e)}")
                return ""
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported"""
        ext = Path(filename).suffix.lower()
        return ext in self.supported_formats
    
    def get_preview(self, text: str, max_length: int = 500) -> str:
        """Get preview snippet from extracted text"""
        if not text:
            return ""
        
        text = text.strip()
        if len(text) <= max_length:
            return text
        
        # Try to break at sentence or word boundary
        preview = text[:max_length]
        last_period = preview.rfind('.')
        last_space = preview.rfind(' ')
        
        if last_period > max_length * 0.7:
            return preview[:last_period + 1]
        elif last_space > max_length * 0.7:
            return preview[:last_space] + '...'
        else:
            return preview + '...'


# Singleton instance
text_extractor = TextExtractor()
