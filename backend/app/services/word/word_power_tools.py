"""
Word document power tools
Uses python-docx for DOCX processing. Async-friendly executor interface.
"""
import io
import logging
from typing import Tuple, Optional

from docx import Document

logger = logging.getLogger(__name__)


class WordPowerToolsExecutor:
    def __init__(self):
        pass

    async def extract_text(self, content: bytes) -> Tuple[Optional[str], Optional[str]]:
        try:
            doc = Document(io.BytesIO(content))
            texts = []
            for para in doc.paragraphs:
                texts.append(para.text)
            return "\n".join(texts), None
        except Exception as e:
            logger.error("Word extract_text error: %s", e)
            return None, str(e)

    async def convert_to_pdf(self, content: bytes) -> Tuple[Optional[bytes], Optional[str]]:
        try:
            import io
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            
            doc = Document(io.BytesIO(content))
            pdf_buffer = io.BytesIO()
            
            pdf = SimpleDocTemplate(
                pdf_buffer,
                pagesize=letter,
                rightMargin=54,
                leftMargin=54,
                topMargin=54,
                bottomMargin=54
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    story.append(Spacer(1, 10))
                    continue
                
                # Simple style mapping
                style = styles['Normal']
                if para.style.name.startswith('Heading 1'):
                    style = styles['Heading1']
                elif para.style.name.startswith('Heading 2'):
                    style = styles['Heading2']
                elif para.style.name.startswith('Heading 3'):
                    style = styles['Heading3']
                elif para.style.name.startswith('Title'):
                    style = styles['Title']
                
                # Build Paragraph
                story.append(Paragraph(text, style))
                story.append(Spacer(1, 6))
                
            pdf.build(story)
            return pdf_buffer.getvalue(), None
        except Exception as e:
            logger.error("Word convert_to_pdf error: %s", e)
            return None, str(e)

    async def merge_documents(self, contents: list[bytes]) -> Tuple[Optional[bytes], Optional[str]]:
        try:
            # Simple approach: create new document and append paragraphs
            merged = Document()
            for i, c in enumerate(contents):
                doc = Document(io.BytesIO(c))
                if i > 0:
                    merged.add_page_break()
                for para in doc.paragraphs:
                    merged.add_paragraph(para.text)
            buf = io.BytesIO()
            merged.save(buf)
            return buf.getvalue(), None
        except Exception as e:
            logger.error("Word merge error: %s", e)
            return None, str(e)

    async def replace_text(self, content: bytes, find_text: str, replace_text: str) -> Tuple[Optional[bytes], Optional[str]]:
        try:
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                if find_text in para.text:
                    para.text = para.text.replace(find_text, replace_text)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue(), None
        except Exception as e:
            logger.error("Word replace_text error: %s", e)
            return None, str(e)

    async def add_watermark(self, content: bytes, watermark_text: str) -> Tuple[Optional[bytes], Optional[str]]:
        # Placeholder: python-docx does not support complex watermarks easily
        return None, "Watermarking DOCX not implemented (requires advanced processing)."


_executor: WordPowerToolsExecutor = None


def get_word_executor() -> WordPowerToolsExecutor:
    global _executor
    if _executor is None:
        _executor = WordPowerToolsExecutor()
    return _executor
